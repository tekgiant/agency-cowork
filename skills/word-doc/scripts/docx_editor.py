"""
Word Document Editor.

Primary tool for inspecting and modifying existing .docx files.
Workflow: download from SharePoint → inspect → edit → upload back.

Usage:
    python -m scripts.docx_editor --input doc.docx --action inspect
    python -m scripts.docx_editor --input doc.docx --action extract-text
    python -m scripts.docx_editor --input doc.docx --action find-replace --find "OLD" --replace "NEW"
    python -m scripts.docx_editor --input doc.docx --action update-paragraph --index 3 --text "New text"
    python -m scripts.docx_editor --input doc.docx --action insert-paragraph --index 5 --text "Inserted" --style "Heading 2"
    python -m scripts.docx_editor --input doc.docx --action update-table-cell --table 0 --row 1 --col 2 --text "Updated"
    python -m scripts.docx_editor --input doc.docx --action add-table --headers-json '["A","B"]' --rows-json '[["1","2"]]'
    python -m scripts.docx_editor --input doc.docx --action update-header --text "Header text"
    python -m scripts.docx_editor --input doc.docx --action update-footer --text "Footer text"
    python -m scripts.docx_editor --input doc.docx --action batch --ops-json '[...]'
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))
from office_common.config import load_config, OfficeConfig


# ─── Inspect ──────────────────────────────────────────────────────────────────

def inspect_document(doc: Document, max_preview: int = 30) -> dict:
    """Comprehensive inspection of a Word document."""
    info = {
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "section_count": len(doc.sections),
    }

    styles_used = set()
    for p in doc.paragraphs:
        if p.style and p.style.name:
            styles_used.add(p.style.name)
    info["styles_used"] = sorted(styles_used)

    paragraphs = []
    for i, p in enumerate(doc.paragraphs):
        if i >= max_preview and i < len(doc.paragraphs) - 5:
            if i == max_preview:
                paragraphs.append({
                    "para_index": "...",
                    "note": f"({len(doc.paragraphs) - max_preview - 5} paragraphs omitted)",
                })
            continue

        p_info = {
            "para_index": i,
            "style": p.style.name if p.style else None,
            "text": p.text[:200] + ("..." if len(p.text) > 200 else ""),
        }

        if p.alignment:
            p_info["alignment"] = str(p.alignment)

        if len(p.runs) > 0 and len(p.text) <= 200:
            runs = []
            for r in p.runs:
                run_info = {"text": r.text}
                if r.bold:
                    run_info["bold"] = True
                if r.italic:
                    run_info["italic"] = True
                if r.underline:
                    run_info["underline"] = True
                if r.font.size:
                    run_info["size_pt"] = round(r.font.size / 12700, 1)
                runs.append(run_info)
            p_info["runs"] = runs

        paragraphs.append(p_info)
    info["paragraphs"] = paragraphs

    tables = []
    for ti, table in enumerate(doc.tables):
        t_info = {
            "table_index": ti,
            "rows": len(table.rows),
            "cols": len(table.columns),
        }
        if table.rows:
            t_info["headers"] = [cell.text[:50] for cell in table.rows[0].cells]
        preview = []
        for ri in range(1, min(4, len(table.rows))):
            preview.append([cell.text[:50] for cell in table.rows[ri].cells])
        t_info["preview_rows"] = preview
        tables.append(t_info)
    info["tables"] = tables

    for si, section in enumerate(doc.sections):
        sec_info = {"section_index": si}
        try:
            if section.header and section.header.paragraphs:
                sec_info["header"] = " | ".join(
                    p.text for p in section.header.paragraphs if p.text.strip())
        except Exception:
            pass
        try:
            if section.footer and section.footer.paragraphs:
                sec_info["footer"] = " | ".join(
                    p.text for p in section.footer.paragraphs if p.text.strip())
        except Exception:
            pass
        sec_info["page_width_inches"] = round(section.page_width / 914400, 2)
        sec_info["page_height_inches"] = round(section.page_height / 914400, 2)

        if si == 0:
            info["primary_section"] = sec_info
        elif "sections" not in info:
            info["sections"] = [sec_info]
        else:
            info["sections"].append(sec_info)

    return info


# ─── Edit Operations ──────────────────────────────────────────────────────────

def find_replace(doc: Document, find: str, replace: str) -> int:
    """Find and replace text. Preserves formatting."""
    count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if find in run.text:
                run.text = run.text.replace(find, replace)
                count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if find in run.text:
                            run.text = run.text.replace(find, replace)
                            count += 1
    return count


def update_paragraph(doc: Document, index: int, text: str,
                     style: Optional[str] = None) -> bool:
    """Update text of a paragraph. Preserves run formatting."""
    if index >= len(doc.paragraphs):
        return False
    para = doc.paragraphs[index]
    if style:
        para.style = doc.styles[style]
    if para.runs:
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.text = text
    return True


def insert_paragraph(doc: Document, index: int, text: str,
                     style: Optional[str] = None) -> bool:
    """Insert a new paragraph before the given index."""
    if index >= len(doc.paragraphs):
        doc.add_paragraph(text, style=style)
    else:
        target = doc.paragraphs[index]
        new_p = target._element.addprevious(
            target._element.makeelement(qn('w:p'), {}))
        from docx.text.paragraph import Paragraph
        para = Paragraph(new_p, target._element.getparent())
        para.add_run(text)
        if style:
            para.style = doc.styles[style]
    return True


def delete_paragraph(doc: Document, index: int) -> bool:
    """Delete a paragraph."""
    if index >= len(doc.paragraphs):
        return False
    p = doc.paragraphs[index]._element
    p.getparent().remove(p)
    return True


def update_table_cell(doc: Document, table_idx: int, row: int, col: int,
                      text: str) -> bool:
    """Update a table cell. Preserves formatting."""
    if table_idx >= len(doc.tables):
        return False
    table = doc.tables[table_idx]
    if row >= len(table.rows) or col >= len(table.columns):
        return False
    cell = table.cell(row, col)
    if cell.paragraphs and cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = text
        for r in cell.paragraphs[0].runs[1:]:
            r.text = ""
    else:
        cell.text = text
    return True


def add_table_row(doc: Document, table_idx: int, values: list[str]) -> bool:
    """Append a row to a table."""
    if table_idx >= len(doc.tables):
        return False
    table = doc.tables[table_idx]
    row = table.add_row()
    for i, val in enumerate(values):
        if i < len(row.cells):
            row.cells[i].text = str(val)
    return True


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    """Add a new table to end of document."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            if i < len(row.cells):
                row.cells[i].text = str(val)


def update_header(doc: Document, text: str, section_idx: int = 0) -> bool:
    if section_idx >= len(doc.sections):
        return False
    section = doc.sections[section_idx]
    header = section.header
    header.is_linked_to_previous = False
    if header.paragraphs:
        header.paragraphs[0].text = text
    else:
        header.add_paragraph(text)
    return True


def update_footer(doc: Document, text: str, section_idx: int = 0) -> bool:
    if section_idx >= len(doc.sections):
        return False
    section = doc.sections[section_idx]
    footer = section.footer
    footer.is_linked_to_previous = False
    if footer.paragraphs:
        footer.paragraphs[0].text = text
    else:
        footer.add_paragraph(text)
    return True


def extract_text(doc: Document) -> str:
    parts = []
    for para in doc.paragraphs:
        parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


# ─── Batch Operations ─────────────────────────────────────────────────────────

def batch_ops(doc: Document, operations: list[dict],
              config: Optional[OfficeConfig] = None) -> list[dict]:
    """Execute a batch of edit operations."""
    results = []
    for op in operations:
        action = op.get("action")
        try:
            if action == "find-replace":
                count = find_replace(doc, op["find"], op["replace"])
                results.append({"action": action, "ok": True, "replacements": count})
            elif action == "update-paragraph":
                ok = update_paragraph(doc, op["index"], op["text"], op.get("style"))
                results.append({"action": action, "ok": ok})
            elif action == "insert-paragraph":
                ok = insert_paragraph(doc, op["index"], op["text"], op.get("style"))
                results.append({"action": action, "ok": ok})
            elif action == "delete-paragraph":
                ok = delete_paragraph(doc, op["index"])
                results.append({"action": action, "ok": ok})
            elif action == "update-table-cell":
                ok = update_table_cell(doc, op["table"], op["row"], op["col"], op["text"])
                results.append({"action": action, "ok": ok})
            elif action == "add-table-row":
                ok = add_table_row(doc, op["table"], op["values"])
                results.append({"action": action, "ok": ok})
            elif action == "add-table":
                add_table(doc, op["headers"], op["rows"])
                results.append({"action": action, "ok": True})
            elif action == "update-header":
                ok = update_header(doc, op["text"], op.get("section", 0))
                results.append({"action": action, "ok": ok})
            elif action == "update-footer":
                ok = update_footer(doc, op["text"], op.get("section", 0))
                results.append({"action": action, "ok": ok})
            else:
                results.append({"action": action, "ok": False, "error": f"Unknown: {action}"})
        except Exception as e:
            results.append({"action": action, "ok": False, "error": str(e)})
    return results


# ─── CLI ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Inspect and edit Word documents")
    parser.add_argument("--input", "-i", required=True, help="Input .docx file")
    parser.add_argument("--output", "-o", help="Output file (default: overwrites input)")
    parser.add_argument("--action", required=True,
                        choices=["inspect", "extract-text", "find-replace",
                                 "update-paragraph", "insert-paragraph",
                                 "delete-paragraph", "update-table-cell",
                                 "add-table-row", "add-table",
                                 "update-header", "update-footer", "batch"],
                        help="Action to perform")
    parser.add_argument("--index", type=int, help="Paragraph index")
    parser.add_argument("--text", help="Text content")
    parser.add_argument("--style", help="Paragraph style name")
    parser.add_argument("--find", help="Text to find")
    parser.add_argument("--replace", help="Replacement text")
    parser.add_argument("--table", type=int, help="Table index")
    parser.add_argument("--row", type=int, help="Table row")
    parser.add_argument("--col", type=int, help="Table column")
    parser.add_argument("--section", type=int, default=0, help="Section index")
    parser.add_argument("--headers-json", help="Table headers JSON")
    parser.add_argument("--rows-json", help="Table rows JSON")
    parser.add_argument("--values-json", help="Row values JSON")
    parser.add_argument("--ops-json", help="Batch operations JSON")
    parser.add_argument("--max-preview", type=int, default=30, help="Max paragraphs to preview")
    parser.add_argument("--config", help="Path to agentconfig.json")
    args = parser.parse_args()

    doc = Document(args.input)
    output_path = args.output or args.input

    if args.action == "inspect":
        result = inspect_document(doc, args.max_preview)
        print(json.dumps(result, indent=2, default=str))
        return

    elif args.action == "extract-text":
        print(extract_text(doc))
        return

    elif args.action == "find-replace":
        if not args.find:
            print("Error: --find required", file=sys.stderr); sys.exit(1)
        count = find_replace(doc, args.find, args.replace or "")
        doc.save(output_path)
        print(json.dumps({"status": "ok", "replacements": count}))

    elif args.action == "update-paragraph":
        if args.index is None or not args.text:
            print("Error: --index and --text required", file=sys.stderr); sys.exit(1)
        ok = update_paragraph(doc, args.index, args.text, args.style)
        doc.save(output_path)
        print(json.dumps({"status": "ok" if ok else "failed"}))

    elif args.action == "insert-paragraph":
        if args.index is None or not args.text:
            print("Error: --index and --text required", file=sys.stderr); sys.exit(1)
        ok = insert_paragraph(doc, args.index, args.text, args.style)
        doc.save(output_path)
        print(json.dumps({"status": "ok" if ok else "failed"}))

    elif args.action == "delete-paragraph":
        if args.index is None:
            print("Error: --index required", file=sys.stderr); sys.exit(1)
        ok = delete_paragraph(doc, args.index)
        doc.save(output_path)
        print(json.dumps({"status": "ok" if ok else "failed"}))

    elif args.action == "update-table-cell":
        if None in (args.table, args.row, args.col) or not args.text:
            print("Error: --table, --row, --col, --text required", file=sys.stderr); sys.exit(1)
        ok = update_table_cell(doc, args.table, args.row, args.col, args.text)
        doc.save(output_path)
        print(json.dumps({"status": "ok" if ok else "failed"}))

    elif args.action == "add-table-row":
        if args.table is None or not args.values_json:
            print("Error: --table and --values-json required", file=sys.stderr); sys.exit(1)
        values = json.loads(args.values_json)
        ok = add_table_row(doc, args.table, values)
        doc.save(output_path)
        print(json.dumps({"status": "ok" if ok else "failed"}))

    elif args.action == "add-table":
        if not args.headers_json or not args.rows_json:
            print("Error: --headers-json and --rows-json required", file=sys.stderr); sys.exit(1)
        add_table(doc, json.loads(args.headers_json), json.loads(args.rows_json))
        doc.save(output_path)
        print(json.dumps({"status": "ok"}))

    elif args.action == "update-header":
        if not args.text:
            print("Error: --text required", file=sys.stderr); sys.exit(1)
        ok = update_header(doc, args.text, args.section)
        doc.save(output_path)
        print(json.dumps({"status": "ok" if ok else "failed"}))

    elif args.action == "update-footer":
        if not args.text:
            print("Error: --text required", file=sys.stderr); sys.exit(1)
        ok = update_footer(doc, args.text, args.section)
        doc.save(output_path)
        print(json.dumps({"status": "ok" if ok else "failed"}))

    elif args.action == "batch":
        if not args.ops_json:
            print("Error: --ops-json required", file=sys.stderr); sys.exit(1)
        config = load_config(args.config)
        results = batch_ops(doc, json.loads(args.ops_json), config)
        doc.save(output_path)
        print(json.dumps({"status": "ok", "results": results}))


if __name__ == "__main__":
    main()
